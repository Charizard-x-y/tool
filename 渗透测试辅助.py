"""
渗透测试工具
作者：Charizard_xy
功能：辅助渗透测试
"""
# 仅做学习使用，请勿用于非法用途
import tkinter as tk
from ttkbootstrap import Style, Notebook, Frame, Button, Label, Entry, Text, Combobox
from ttkbootstrap.constants import *
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import re
import os
from tkinter import filedialog, messagebox
from datetime import datetime

class PenTestTool:
    def __init__(self, master):
        self.master = master
        self.master.title("渗透测试工具")
        self.style = Style(theme='darkly')
        
        # 初始化数据存储
        self.data = {
            "info_gathering": {
                "target_url": "",
                "ip_address": "",
                "domain": "",
                "open_ports": "",
                "subdomains": "",
                "whois_info": "",
                "server_type": ""
            },
            "vuln_scan": {
                "scan_type": "全面扫描",
                "cve_list": "",
                "cvss_scores": "",
                "vuln_descriptions": "",
                "proof_of_concept": ""
            },
            "exploitation": {
                "exploit_module": "",
                "payload_config": "",
                "session_management": "",
                "obtained_privileges": "",
                "exploit_steps": ""
            },
            "post_exploit": {
                "privilege_escalation": "",
                "lateral_movement": "",
                "data_exfiltration": "",
                "persistence_method": "",
                "network_topology": ""
            },
            "reporting": {
                "title": "渗透测试报告",
                "author": "",
                "client": "",
                "notes": ""
            }
        }
        
        # 审计日志系统
        self.audit_log = []
        self.current_project = None
        self.init_audit_log()
        
        # 创建界面组件
        self.create_widgets()
        self.log_window = None

    def init_audit_log(self):
        """初始化审计日志目录"""
        log_dir = "audit_logs"
        os.makedirs(log_dir, exist_ok=True)

    def create_widgets(self):
        # 主容器
        main_frame = Frame(self.master)
        main_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        # 创建选项卡
        self.notebook = Notebook(main_frame)
        tabs = [
            ("信息收集", self.create_info_gathering_tab),
            ("漏洞扫描", self.create_vuln_scan_tab),
            ("漏洞利用", self.create_exploitation_tab),
            ("后渗透", self.create_post_exploit_tab),
            ("报告生成", self.create_reporting_tab)
        ]
        
        self.tab_frames = {}
        for text, command in tabs:
            frame = Frame(self.notebook)
            self.tab_frames[text] = frame
            self.notebook.add(frame, text=text)
            command(frame)
        
        self.notebook.pack(fill=BOTH, expand=True)
        
        # 主控制按钮
        control_frame = Frame(main_frame)
        Button(control_frame, text="生成报告", command=self.generate_report).pack(side=LEFT, padx=5)
        Button(control_frame, text="清空所有数据", command=self.confirm_clear_all).pack(side=LEFT, padx=5)
        Button(control_frame, text="审计日志", command=self.show_audit_log).pack(side=LEFT, padx=5)
        control_frame.pack(pady=10)

    def create_info_gathering_tab(self, parent):
        frame = Frame(parent)
        frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        fields = [
            ("目标URL*", "target_url", "entry", r'^https?://\S+'),
            ("IP地址*", "ip_address", "entry", r'^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}$'),
            ("域名", "domain", "entry", None),
            ("开放端口*", "open_ports", "entry", r'^(\d+,)*\d+$'),
            ("子域名", "subdomains", "text", None),
            ("WHOIS信息", "whois_info", "text", None),
            ("服务器类型", "server_type", "entry", None)
        ]
        
        self.create_section_fields(frame, fields, "info_gathering")

    def create_vuln_scan_tab(self, parent):
        frame = Frame(parent)
        frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        fields = [
            ("扫描类型", "scan_type", "combobox", ["全面扫描", "漏洞扫描", "合规检查"]),
            ("CVE列表*", "cve_list", "text", r'^CVE-\d{4}-\d+'),
            ("CVSS评分", "cvss_scores", "text", r'^\d+\.\d+'),
            ("漏洞描述*", "vuln_descriptions", "text", None),
            ("验证POC", "proof_of_concept", "text", None)
        ]
        
        self.create_section_fields(frame, fields, "vuln_scan")

    def create_exploitation_tab(self, parent):
        frame = Frame(parent)
        frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        fields = [
            ("利用模块*", "exploit_module", "entry", None),
            ("Payload配置*", "payload_config", "text", None),
            ("会话管理", "session_management", "text", None),
            ("获得权限", "obtained_privileges", "combobox", ["user", "admin", "system"]),
            ("利用步骤*", "exploit_steps", "text", None)
        ]
        
        self.create_section_fields(frame, fields, "exploitation")

    def create_post_exploit_tab(self, parent):
        frame = Frame(parent)
        frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        fields = [
            ("权限提升", "privilege_escalation", "text", None),
            ("横向移动", "lateral_movement", "text", None),
            ("数据提取", "data_exfiltration", "text", None),
            ("持久化方法", "persistence_method", "combobox", ["注册表", "计划任务", "服务"]),
            ("网络拓扑", "network_topology", "text", None)
        ]
        
        self.create_section_fields(frame, fields, "post_exploit")

    def create_reporting_tab(self, parent):
        frame = Frame(parent)
        frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        fields = [
            ("报告标题*", "title", "entry", None),
            ("测试人员*", "author", "entry", None),
            ("客户名称*", "client", "entry", None),
            ("备注说明", "notes", "text", None)
        ]
        
        self.create_section_fields(frame, fields, "reporting")

    def create_section_fields(self, parent, fields, section):
        for label_text, field_name, field_type, validation in fields:
            Label(parent, text=label_text).pack(anchor=NW)
            
            if field_type == "entry":
                widget = Entry(parent)
                widget.pack(fill=X, pady=2)
            elif field_type == "text":
                widget = Text(parent, height=4)
                widget.pack(fill=X, pady=2)
            elif field_type == "combobox":
                widget = Combobox(parent, values=validation)
                widget.pack(fill=X, pady=2)
            
            setattr(self, f"{section}_{field_name}", widget)
            setattr(self, f"{section}_{field_name}_valid", validation)

        btn_frame = Frame(parent)
        Button(btn_frame, text="保存数据", command=lambda: self.save_data(section)).pack(side=LEFT, padx=5)
        Button(btn_frame, text="清空当前", command=lambda: self.confirm_clear_section(section)).pack(side=LEFT, padx=5)
        btn_frame.pack(pady=10)

    def validate_input(self, value, pattern):
        if not pattern:
            return True
        if isinstance(pattern, list):
            return value in pattern
        return re.match(pattern, value) is not None

    def save_data(self, section):
        try:
            data = {}
            validation_errors = []
            
            for widget in dir(self):
                if widget.startswith(f"{section}_") and not widget.endswith("_valid"):
                    field_name = widget.split("_", 2)[-1]
                    validation_pattern = getattr(self, f"{section}_{field_name}_valid", None)
                    
                    entry = getattr(self, widget)
                    if isinstance(entry, Entry):
                        value = entry.get()
                    elif isinstance(entry, Text):
                        value = entry.get("1.0", END).strip()
                    elif isinstance(entry, Combobox):
                        value = entry.get()
                    
                    label = next((f[0] for f in self.get_section_fields(section) if f[1] == field_name), "")
                    if label.endswith("*") and not value:
                        validation_errors.append(f"{label[:-1]} 是必填字段")
                    elif validation_pattern and not self.validate_input(value, validation_pattern):
                        validation_errors.append(f"{label} 格式无效")
                    
                    data[field_name] = value
            
            if validation_errors:
                self.show_message("验证错误", "\n".join(validation_errors), DANGER)
                self.log_action("数据验证失败", section, "\n".join(validation_errors))
                return
            
            self.data[section] = data
            self.update_tab_color(section, SUCCESS)
            self.show_message("保存成功", "数据已保存", SUCCESS)
            self.log_action("数据保存", section, f"{section}模块数据更新")
            
        except Exception as e:
            self.log_action("保存失败", section, str(e))
            messagebox.showerror("错误", f"保存失败: {str(e)}")

    def update_tab_color(self, section, color):
        tab_names = {
            "info_gathering": 0,
            "vuln_scan": 1,
            "exploitation": 2,
            "post_exploit": 3,
            "reporting": 4
        }
        self.notebook.tab(tab_names[section], style=color)

    def confirm_clear_section(self, section):
        if messagebox.askyesno("确认", "确定要清空当前模块数据吗？"):
            self.clear_section(section)
            self.log_action("数据清除", section, "清空当前模块数据")

    def clear_section(self, section):
        for widget in dir(self):
            if widget.startswith(f"{section}_"):
                entry = getattr(self, widget)
                if isinstance(entry, Entry):
                    entry.delete(0, END)
                elif isinstance(entry, Text):
                    entry.delete("1.0", END)
                elif isinstance(entry, Combobox):
                    entry.set('')
        self.data[section] = {}
        self.update_tab_color(section, DANGER)

    def confirm_clear_all(self):
        if messagebox.askyesno("确认", "确定要清空所有数据吗？"):
            self.clear_all_data()
            self.log_action("全局操作", details="清空所有数据")

    def clear_all_data(self):
        for section in self.data:
            self.clear_section(section)
        self.current_project = None

    def generate_report(self):
        try:
            if not self.validate_all_sections():
                return
            
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")]
            )
            if not filename:
                return
            
            self.generate_word_report(filename)
            self.show_message("报告生成", f"报告已保存为: {filename}", SUCCESS)
            self.log_action("报告生成", details=f"生成报告: {filename}")
            
        except Exception as e:
            self.log_action("报告生成失败", details=str(e))
            messagebox.showerror("错误", f"报告生成失败: {str(e)}")

    def generate_word_report(self, filename):
        doc = Document()
        
        # 封面
        self.add_cover_page(doc)
        
        # 各章节
        sections = [
            ("信息收集", self.data["info_gathering"]),
            ("漏洞扫描", self.data["vuln_scan"]),
            ("漏洞利用", self.data["exploitation"]),
            ("后渗透", self.data["post_exploit"]),
            ("测试报告", self.data["reporting"])
        ]
        
        for title, data in sections:
            self.add_section(doc, title, data)
        
        # 审计日志
        self.add_audit_log_section(doc)
        
        doc.save(filename)

    def add_cover_page(self, doc):
        title = doc.add_paragraph()
        title_run = title.add_run(self.data["reporting"]["title"])
        title_run.font.size = Pt(28)
        title_run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        details = [
            ("客户名称", self.data["reporting"]["client"]),
            ("测试人员", self.data["reporting"]["author"]),
            ("报告日期", datetime.now().strftime("%Y-%m-%d"))
        ]
        
        for label, value in details:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)
        
        doc.add_page_break()

    def add_section(self, doc, title, data):
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "LightShading"
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "项目"
        hdr_cells[1].text = "内容"
        
        for key, value in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)
        
        doc.add_paragraph()

    def add_audit_log_section(self, doc):
        doc.add_heading("审计日志", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "LightShading"
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '时间'
        hdr_cells[1].text = '操作'
        hdr_cells[2].text = '模块'
        hdr_cells[3].text = '详细信息'
        
        for entry in self.audit_log[-50:]:  # 显示最近50条日志
            row_cells = table.add_row().cells
            row_cells[0].text = entry['timestamp']
            row_cells[1].text = entry['action']
            row_cells[2].text = entry.get('module', '')
            row_cells[3].text = str(entry.get('details', ''))[:100]

    def validate_all_sections(self):
        missing_sections = [name for name, data in self.data.items() if not data]
        if missing_sections:
            self.show_message("验证失败", "以下模块未完成：\n" + "\n".join(missing_sections), DANGER)
            return False
        return True

    def log_action(self, action, module=None, details=None):
        log_entry = {
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "action": action,
            "module": module,
            "details": details
        }
        self.audit_log.append(log_entry)
        self.save_audit_log(log_entry)

    def save_audit_log(self, entry):
        log_file = os.path.join("audit_logs", f"audit_{datetime.now().strftime('%Y%m')}.log")
        with open(log_file, 'a', encoding='utf-8') as f:
            f.write(json.dumps(entry, ensure_ascii=False) + '\n')

    def show_audit_log(self):
        if self.log_window and self.log_window.winfo_exists():
            self.log_window.lift()
            return
        
        self.log_window = tk.Toplevel(self.master)
        self.log_window.title("审计日志查看器")
        self.log_window.geometry("1000x600")
        
        main_frame = Frame(self.log_window)
        main_frame.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        # 工具栏
        toolbar = Frame(main_frame)
        Button(toolbar, text="导出日志", command=self.export_audit_log).pack(side=LEFT)
        Button(toolbar, text="刷新", command=self.refresh_log_view).pack(side=LEFT)
        toolbar.pack(fill=X)
        
        # 日志显示
        log_text = Text(main_frame, wrap=WORD)
        scrollbar = tk.Scrollbar(main_frame, command=log_text.yview)
        log_text.configure(yscrollcommand=scrollbar.set)
        
        log_text.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.load_log_to_text(log_text)

    def load_log_to_text(self, text_widget):
        text_widget.delete(1.0, END)
        for entry in reversed(self.audit_log[-100:]):  # 显示最近100条
            text_widget.insert(END, f"[{entry['timestamp']}] {entry['action']}\n")
            text_widget.insert(END, f"模块: {entry.get('module', '全局')}\n")
            text_widget.insert(END, f"详情: {entry.get('details', '')}\n")
            text_widget.insert(END, "-"*80 + "\n")

    def export_audit_log(self):
        filename = filedialog.asksaveasfilename(
            defaultextension=".log",
            filetypes=[("Log Files", "*.log"), ("All Files", "*.*")]
        )
        if filename:
            with open(filename, 'w', encoding='utf-8') as f:
                for entry in self.audit_log:
                    f.write(json.dumps(entry, ensure_ascii=False) + '\n')
            messagebox.showinfo("导出成功", f"日志已导出到：{filename}")

    def refresh_log_view(self):
        if self.log_window:
            for widget in self.log_window.winfo_children():
                if isinstance(widget, Text):
                    self.load_log_to_text(widget)
                    break

    def save_project(self):
        try:
            filename = filedialog.asksaveasfilename(
                defaultextension=".json",
                filetypes=[("JSON Files", "*.json")]
            )
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    json.dump(self.data, f, ensure_ascii=False)
                self.current_project = filename
                self.log_action("项目保存", details=f"保存到：{filename}")
                messagebox.showinfo("保存成功", "项目文件已保存")
        except Exception as e:
            self.log_action("项目保存失败", details=str(e))
            messagebox.showerror("错误", f"保存失败：{str(e)}")

    def load_project(self):
        try:
            filename = filedialog.askopenfilename(
                filetypes=[("JSON Files", "*.json")]
            )
            if filename:
                with open(filename, 'r', encoding='utf-8') as f:
                    self.data = json.load(f)
                self.update_ui_from_data()
                self.current_project = filename
                self.log_action("项目加载", details=f"加载项目：{filename}")
                messagebox.showinfo("加载成功", "项目文件已加载")
        except Exception as e:
            self.log_action("项目加载失败", details=str(e))
            messagebox.showerror("错误", f"加载失败：{str(e)}")

    def update_ui_from_data(self):
        for section in self.data:
            for field, value in self.data[section].items():
                widget = getattr(self, f"{section}_{field}", None)
                if widget:
                    if isinstance(widget, Entry):
                        widget.delete(0, END)
                        widget.insert(0, value)
                    elif isinstance(widget, Text):
                        widget.delete("1.0", END)
                        widget.insert("1.0", value)
                    elif isinstance(widget, Combobox):
                        widget.set(value)
            self.update_tab_color(section, SUCCESS if self.data[section] else DANGER)

    def show_message(self, title, message, style):
        top = tk.Toplevel(self.master)
        top.title(title)
        Label(top, text=message, style=style).pack(padx=20, pady=20)
        Button(top, text="确定", command=top.destroy).pack(pady=10)

if __name__ == "__main__":
    root = tk.Tk()
    app = PenTestTool(root)
    root.geometry("1200x800")
    
    # 创建菜单系统
    menu_bar = tk.Menu(root)
    
    # 文件菜单
    file_menu = tk.Menu(menu_bar, tearoff=0)
    file_menu.add_command(label="新建项目", command=app.confirm_clear_all)
    file_menu.add_command(label="打开项目", command=app.load_project)
    file_menu.add_command(label="保存项目", command=app.save_project)
    file_menu.add_separator()
    file_menu.add_command(label="退出", command=root.quit)
    menu_bar.add_cascade(label="文件", menu=file_menu)
    
    # 日志菜单
    log_menu = tk.Menu(menu_bar, tearoff=0)
    log_menu.add_command(label="查看日志", command=app.show_audit_log)
    log_menu.add_command(label="导出日志", command=app.export_audit_log)
    menu_bar.add_cascade(label="审计日志", menu=log_menu)
    
    root.config(menu=menu_bar)
    root.mainloop()
# 仅做学习使用，请勿用于非法用途